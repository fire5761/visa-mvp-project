import os
import re
import json
import shutil
import subprocess
import uuid
import hashlib
from pathlib import Path
from datetime import datetime, date, timezone

import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, field_validator, model_validator, EmailStr
from starlette.background import BackgroundTask
from google.cloud import bigquery
from google.cloud import firestore
import chromadb

# --- [Vertex AI 라이브러리 추가] ---
import vertexai
from vertexai.generative_models import GenerativeModel
from vertexai.language_models import TextEmbeddingModel, TextEmbeddingInput

import sendgrid
from sendgrid.helpers.mail import Mail
from fastapi import Header
from datetime import timedelta

# 1. GCP 서비스 계정 인증 (로컬 환경 변수 주입)
#os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "assets/gcp-key.json"

# 2. Vertex AI 초기화
PROJECT_ID = "proj-aj29-211200020328"
LOCATION = "global"                  # Gemini 생성형 모델용 (gemini-3.1-flash-lite는 global만 지원)
EMBEDDING_LOCATION = "us-central1"   # 임베딩 모델용 (gemini-embedding-001은 global 미지원, 리전 필요)

# 5-1. 임베딩 모델을 먼저 리전(us-central1) 컨텍스트에서 생성
vertexai.init(project=PROJECT_ID, location=EMBEDDING_LOCATION)
EMBEDDING_MODEL_NAME = "gemini-embedding-001"
embedding_model = TextEmbeddingModel.from_pretrained(EMBEDDING_MODEL_NAME)

# 2-1. 이후 global 컨텍스트로 전환해서 Gemini 생성형 모델 생성
vertexai.init(project=PROJECT_ID, location=LOCATION)
model = GenerativeModel(model_name="gemini-3.1-flash-lite")

# 3. BigQuery 초기화 (F-03 서류 요구사항 지식베이스 캐싱용)
BQ_DATASET = "visa_mvp"
BQ_TABLE = "document_requirements"
bq_client = bigquery.Client(project=PROJECT_ID)

# 4. Firestore 초기화 (F-01 최소 알림정보 저장용)
fs_client = firestore.Client(project=PROJECT_ID)
NOTIFICATION_COLLECTION = "notification_subscriptions"

# 5-2. RAG(ChromaDB) 초기화 — build_rag_index.py로 미리 인덱싱된 chroma_db를 읽기만 함
CHROMA_PERSIST_DIR = "chroma_db"
CHROMA_COLLECTION_NAME = "visa_manual"

chroma_client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)
rag_collection = chroma_client.get_collection(CHROMA_COLLECTION_NAME)

app = FastAPI(title="Visa MVP API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # MVP 단계라 전체 허용, 배포 시엔 실제 프론트 도메인으로 좁히는 게 좋음
    allow_methods=["*"],
    allow_headers=["*"],
)


# =========================================================
# 신청구분(10종) 체크박스 위치 — 서식(통합신청서)은 신청구분과 무관하게 공통 1개
# =========================================================
UNIVERSAL_TEMPLATE_PATH = "assets/template_universal.docx"

APPLICATION_TYPE_MAP = {
    "외국인등록": {"row": 4, "col": 0, "target": "[  ] ", "new": "[V] "},
    "체류자격외활동허가": {"row": 4, "col": 12, "target": "[  ]", "new": "[V]", "desired_status_cell": (4, 23)},
    "등록증재발급": {"row": 6, "col": 0, "target": "[  ] ", "new": "[V] "},
    "근무처변경추가허가": {"row": 6, "col": 12, "target": "[  ] ", "new": "[V] "},
    "체류기간연장허가": {"row": 7, "col": 0, "target": "[  ] ", "new": "[V] "},
    "재입국허가": {"row": 7, "col": 12, "target": "[  ] ", "new": "[V] "},
    "체류자격변경허가": {"row": 8, "col": 0, "target": "[  ] ", "new": "[V] ", "desired_status_cell": (8, 4)},
    "체류지변경신고": {"row": 8, "col": 12, "target": "[  ] ", "new": "[V] "},
    "체류자격부여": {"row": 10, "col": 0, "target": "[  ] ", "new": "[V] ", "desired_status_cell": (10, 4)},
    "등록사항변경신고": {"row": 10, "col": 12, "target": "[  ]", "new": "[V]"},
}

DESIRED_STATUS_REQUIRED_TYPES = {"체류자격외활동허가", "체류자격변경허가", "체류자격부여"}

# 인적사항/연락처/여권 등 필드 위치 — 신청구분과 무관하게 공통
FIELD_MAP = {
    "surname": (14, 2),
    "given_name": (14, 12),
    "birth_year": (16, 3), "birth_month": (16, 13), "birth_day": (16, 17),
    "sex_male": {"row": 15, "col": 27, "para": 0, "target": "[ ]", "new": "[V]"},
    "sex_female": {"row": 15, "col": 27, "para": 1, "target": "[ ]", "new": "[V]"},
    "nationality": {"row": 15, "col": 34, "para": 2},
    "arc_comb_row": 17,
    "arc_comb_cols": [7, 9, 13, 14, 17, 19, 22, 24, 26, 27, 29, 30, 33],
    "passport_number": (18, 2),
    "passport_issue_date": (18, 22),
    "passport_expiry_date": (18, 38),
    "address": (19, 2),
    "phone": (20, 6), "mobile_phone": (20, 34),
    "home_country_address": (21, 6), "home_country_phone": (21, 39),
    "email": (27, 27),
    "application_date": (29, 9),
}


# =========================================================
# 데이터 모델
# =========================================================
class UserProfile(BaseModel):
    application_type: str       # "체류기간연장허가" 등 — APPLICATION_TYPE_MAP의 키
    visa_type: str              # "D-2-2" 등 — 첨부서류 판단용
    surname: str
    given_name: str
    date_of_birth: str          # YYYY-MM-DD
    sex: str                    # "M" 또는 "F"
    nationality: str
    arc_number_front: str
    arc_number_back: str
    passport_number: str
    passport_issue_date: str    # YYYY-MM-DD
    passport_expiry_date: str   # YYYY-MM-DD
    visa_expiry_date: str       # YYYY-MM-DD (PDF엔 안 채움, D-Day 알림 계산용)
    address: str                # 대한민국 내 주소
    phone: str
    mobile_phone: str
    home_country_address: str
    home_country_phone: str
    email: EmailStr
    desired_status: str | None = None   # 희망 자격 (예: "D-4-1") — 체류자격외활동허가/변경허가/부여 시에만 필수

    @field_validator("application_type")
    @classmethod
    def validate_application_type(cls, v: str) -> str:
        if v not in APPLICATION_TYPE_MAP:
            raise ValueError(f"지원하지 않는 신청 구분입니다: {v}. 가능한 값: {list(APPLICATION_TYPE_MAP.keys())}")
        return v

    @field_validator("surname", "given_name", "nationality")
    @classmethod
    def uppercase_and_not_empty(cls, v: str) -> str:
        if not v.strip():
            raise ValueError("필수 항목이 비어있습니다.")
        return v.strip().upper()

    @field_validator("sex")
    @classmethod
    def validate_sex(cls, v: str) -> str:
        if v.upper() not in ("M", "F"):
            raise ValueError("성별은 M 또는 F여야 합니다.")
        return v.upper()

    @field_validator("arc_number_front")
    @classmethod
    def validate_arc_front(cls, v: str) -> str:
        if not re.fullmatch(r"\d{6}", v):
            raise ValueError("외국인등록번호 앞자리는 6자리 숫자여야 합니다.")
        return v

    @field_validator("arc_number_back")
    @classmethod
    def validate_arc_back(cls, v: str) -> str:
        if not re.fullmatch(r"\d{7}", v):
            raise ValueError("외국인등록번호 뒷자리는 7자리 숫자여야 합니다.")
        return v

    @field_validator("date_of_birth", "visa_expiry_date", "passport_issue_date", "passport_expiry_date")
    @classmethod
    def validate_date_format(cls, v: str) -> str:
        try:
            datetime.strptime(v, "%Y-%m-%d")
        except ValueError:
            raise ValueError("날짜는 YYYY-MM-DD 형식이어야 합니다.")
        return v

    @field_validator(
        "visa_type", "address", "passport_number",
        "phone", "mobile_phone", "home_country_address", "home_country_phone"
    )
    @classmethod
    def not_empty(cls, v: str) -> str:
        if not v.strip():
            raise ValueError("필수 항목이 비어있습니다.")
        return v.strip()

    @model_validator(mode="after")
    def check_desired_status_required(self):
        if self.application_type in DESIRED_STATUS_REQUIRED_TYPES:
            if not self.desired_status or not self.desired_status.strip():
                raise ValueError(
                    f"'{self.application_type}' 신청 시 희망 자격(예: D-4-1)을 입력해야 합니다."
                )
        return self


class VisaCheckRequest(BaseModel):
    visa_type: str
    application_type: str


@app.get("/")
def read_root():
    return {"message": "Visa MVP Backend is running with Vertex AI!"}


# =========================================================
# F-01: 유저 프로필 등록 + D-Day 알림용 최소정보 저장 (Firestore)
# =========================================================
def save_notification_subscription(email: str, visa_expiry_date: str, application_type: str):
    """D-Day 알림을 위한 최소 정보만 Firestore에 저장.
    개인 식별정보(이름·외국인등록번호 등)는 저장하지 않음.
    문서 ID를 이메일 해시로 지정하여, 같은 사용자가 재등록하면 자동으로 덮어씀(upsert)."""
    user_hash = hashlib.sha256(email.encode("utf-8")).hexdigest()
    doc_ref = fs_client.collection(NOTIFICATION_COLLECTION).document(user_hash)
    doc_ref.set({
        "email": email,
        "visa_expiry_date": visa_expiry_date,
        "application_type": application_type,
        "updated_at": datetime.now(timezone.utc).isoformat(),
    })


@app.post("/api/v1/user/profile")
def register_user_profile(profile: UserProfile):
    try:
        save_notification_subscription(profile.email, profile.visa_expiry_date, profile.application_type)
        return {"status": "success", "profile": profile.model_dump()}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


# =========================================================
# F-03: AI 맞춤형 서류 추천
# 1) BigQuery에 이미 있는 조합이면 즉시 반환(캐시)
# 2) 없으면 RAG(ChromaDB 벡터검색)로 매뉴얼 전체가 아닌 관련 청크만 추출
# 3) 그 청크를 컨텍스트로 Gemini 호출 → 결과를 BigQuery에 저장하여 다음부터는 캐시로 처리
# =========================================================
def query_requirements_from_bigquery(visa_type: str, application_type: str) -> list[str]:
    """BigQuery에 이미 저장된 조합인지 먼저 조회."""
    query = f"""
        SELECT DISTINCT required_document
        FROM `{PROJECT_ID}.{BQ_DATASET}.{BQ_TABLE}`
        WHERE visa_type = @visa_type AND application_type = @application_type
    """
    job_config = bigquery.QueryJobConfig(
        query_parameters=[
            bigquery.ScalarQueryParameter("visa_type", "STRING", visa_type),
            bigquery.ScalarQueryParameter("application_type", "STRING", application_type),
        ]
    )
    result = bq_client.query(query, job_config=job_config).result()
    return [row.required_document for row in result]


def save_requirements_to_bigquery(visa_type: str, application_type: str, documents: list[str]):
    """Gemini가 새로 생성한 결과를 다음 요청부터는 빠르게 쓸 수 있도록 캐싱."""
    if not documents:
        return
    now = datetime.now(timezone.utc).isoformat()
    rows = [
        {
            "visa_type": visa_type,
            "application_type": application_type,
            "required_document": doc,
            "source": "Gemini 실시간 호출 (자동 캐싱)",
            "created_at": now,
        }
        for doc in documents
    ]
    table_ref = f"{PROJECT_ID}.{BQ_DATASET}.{BQ_TABLE}"
    bq_client.insert_rows_json(table_ref, rows)


def retrieve_chunks_for_query(query_text: str, top_k: int = 4) -> list[str]:
    """임의의 자연어 질의로 매뉴얼에서 관련 청크를 검색 (챗봇 + 서류 체크리스트 공용)."""
    query_embedding = embedding_model.get_embeddings(
        [TextEmbeddingInput(text=query_text, task_type="RETRIEVAL_QUERY")]
    )[0].values
    results = rag_collection.query(query_embeddings=[query_embedding], n_results=top_k)
    return results["documents"][0]


def retrieve_relevant_chunks(visa_type: str, application_type: str, top_k: int = 3) -> list[str]:
    """F-03 서류 체크리스트용 (기존 용도 그대로 유지)."""
    query_text = f"{visa_type} {application_type} 필수 제출 서류"
    return retrieve_chunks_for_query(query_text, top_k)


def get_required_documents(visa_type: str, application_type: str) -> dict:
    # 1. BigQuery에 이미 있는 조합이면 바로 반환 (빠름, Gemini/RAG 호출 없음)
    cached = query_requirements_from_bigquery(visa_type, application_type)
    if cached:
        print(f"⚡ BigQuery 캐시 사용: {visa_type} / {application_type} ({len(cached)}건)")
        return {"required_documents": cached, "source": "bigquery_cache"}

    # 2. 없으면 RAG로 관련 청크만 검색 (매뉴얼 전체를 프롬프트에 넣지 않음)
    relevant_chunks = retrieve_relevant_chunks(visa_type, application_type)
    if not relevant_chunks:
        return {"required_documents": []}
    rag_context = "\n".join(relevant_chunks)

    prompt = f"""
    너는 대한민국 출입국 외국인청의 서류 안내 전문가야.
    아래 [검색된 관련 규정]을 100% 참고해서, 아래 두 조건을 모두 만족하는 필수 첨부서류 목록을 추출해.

    [검색된 관련 규정]:
    {rag_context}

    [체류자격(비자 타입)]: {visa_type}
    [신청 구분]: {application_type}

    답변 조건:
    1. 반드시 JSON 형식으로만 반환할 것. (백틱 ```json ... ``` 등 마크다운 빼고 순수 JSON 텍스트만)
    2. JSON 구조는 다음과 같아야 함:
    {{
        "visa_type": "{visa_type}",
        "application_type": "{application_type}",
        "required_documents": ["서류1", "서류2", "서류3"]
    }}
    """

    print(f"🧠 (캐시 없음, RAG 검색됨) {visa_type} / {application_type}에 대한 서류를 Gemini가 분석 중입니다...")
    response = model.generate_content(prompt)

    try:
        clean_text = response.text.replace("```json", "").replace("```", "").strip()
        result = json.loads(clean_text)
        docs = result.get("required_documents", [])
        save_requirements_to_bigquery(visa_type, application_type, docs)  # 다음부터 빨라지도록 캐싱
        return {"required_documents": docs, "source": "gemini_live_with_rag"}
    except Exception:
        return {"required_documents": [], "error": "AI가 JSON 형식을 반환하지 않았습니다.", "raw_text": response.text}


@app.post("/api/v1/visa/checklist")
def get_visa_checklist(req: VisaCheckRequest):
    return get_required_documents(req.visa_type, req.application_type)


# =========================================================
# F-02: 표 셀 기반 안전한 텍스트 삽입 헬퍼
# =========================================================
def set_cell_text_preserving_format(table, row: int, col: int, value: str):
    """cell.text = value 방식은 기존 서식(폰트/크기)을 날려버리므로,
    첫 번째 run의 텍스트만 직접 바꿔서 원본 폰트/크기를 그대로 유지한다."""
    cell = table.cell(row, col)
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = value
        for extra_run in paragraph.runs[1:]:
            extra_run.text = ""
    else:
        paragraph.add_run(value)


def set_paragraph_text_preserving_format(table, row: int, col: int, paragraph_index: int, value: str):
    """국적처럼 '라벨+빈 답변'이 한 셀 안에 여러 문단으로 나뉜 경우,
    특정 문단의 run만 안전하게 교체."""
    cell = table.cell(row, col)
    paragraph = cell.paragraphs[paragraph_index]
    if paragraph.runs:
        paragraph.runs[0].text = value
        for extra in paragraph.runs[1:]:
            extra.text = ""
    else:
        paragraph.add_run(value)


def set_checkbox_preserving_format(table, row: int, col: int, target_text: str, new_text: str):
    """'[  ] ' 텍스트를 가진 run만 정확히 찾아서 교체.
    다른 run(라벨 텍스트 등)의 서식은 건드리지 않아 줄바꿈/레이아웃 밀림이 없다."""
    cell = table.cell(row, col)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.text == target_text:
                run.text = new_text
                return
    raise HTTPException(status_code=500, detail=f"체크박스 대상 텍스트를 찾을 수 없습니다: {target_text!r}")


def set_radio_checkbox_preserving_format(table, row: int, col: int, paragraph_index: int, target_text: str, new_text: str):
    """성별처럼 같은 셀 안에 '[ ]남'/'[ ]여' 두 옵션이 각각 다른 문단에 있는 경우,
    선택된 옵션의 '[ ]' run만 교체."""
    cell = table.cell(row, col)
    paragraph = cell.paragraphs[paragraph_index]
    for run in paragraph.runs:
        if run.text == target_text:
            run.text = new_text
            return
    raise HTTPException(status_code=500, detail=f"체크박스 대상을 찾을 수 없습니다: {target_text!r}")


def set_desired_status_preserving_format(table, row: int, col: int, value: str):
    """'(희망 자격 :        )' 셀에서 콜론+공백+닫는괄호가 한 run에 뭉쳐있으므로,
    그 run 전체를 ': {값} )' 형태로 교체 (콜론/괄호 구조는 유지)."""
    cell = table.cell(row, col)
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        if run.text.startswith(":") and run.text.endswith(")"):
            run.text = f": {value} )"
            return
    raise HTTPException(status_code=500, detail="희망 자격 입력란을 찾을 수 없습니다.")


def prepend_checklist_page(doc, table, checklist: dict):
    """필요서류 안내문을 별도 파일로 만들지 않고,
    신청서 표(table) 바로 앞에 안내문 페이지를 끼워넣는다.
    → LibreOffice 변환을 1번으로 줄이기 위한 성능 최적화."""
    docs = checklist.get("required_documents", [])
    tbl_element = table._element

    def insert_before_table(text: str = "", bold: bool = False, size: int | None = None):
        temp_p = doc.add_paragraph()          # 일단 문서 맨 끝에 추가됨
        temp_p._p.getparent().remove(temp_p._p)  # 원래 위치에서 떼어냄
        tbl_element.addprevious(temp_p._p)        # 표 바로 앞으로 이동
        if text:
            run = temp_p.add_run(text)
            run.bold = bold
            if size:
                run.font.size = Pt(size)
        return temp_p

    insert_before_table("체류허가 신청 - 필요 서류 안내", bold=True, size=16)
    insert_before_table("아래 서류를 준비하여 본 신청서와 함께 제출하세요.")

    if not docs:
        insert_before_table("⚠ 필요 서류를 자동으로 판단하지 못했습니다. 하이코리아 또는 1345(외국인종합안내센터)에 문의하세요.")
    else:
        for i, item in enumerate(docs, 1):
            insert_before_table(f"{i}. {item}")

    insert_before_table("")
    insert_before_table("※ 본 안내는 참고용이며, 최종 제출 서류는 담당 출입국·외국인관서 확인이 필요합니다.")

    # 안내문과 신청서를 별도 페이지로 분리
    page_break_paragraph = insert_before_table("")
    page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)


def convert_docx_to_pdf(docx_path: str, output_dir: str) -> str:
    """LibreOffice headless로 docx -> pdf 변환.
    동시 요청 시 프로필 충돌을 피하기 위해 매 호출마다 격리된 임시 프로필 사용."""
    profile_dir = f"/tmp/lo_profile_{uuid.uuid4().hex}"
    cmd = [
        "soffice", "--headless", "--nologo", "--nofirststartwizard",
        f"-env:UserInstallation=file://{profile_dir}",
        "--convert-to", "pdf",
        "--outdir", output_dir,
        docx_path,
    ]
    result = subprocess.run(cmd, capture_output=True, timeout=30)
    shutil.rmtree(profile_dir, ignore_errors=True)

    if result.returncode != 0:
        raise HTTPException(
            status_code=500,
            detail=f"LibreOffice 변환 실패: {result.stderr.decode(errors='ignore')}"
        )

    pdf_path = str(Path(output_dir) / (Path(docx_path).stem + ".pdf"))
    if not os.path.exists(pdf_path):
        raise HTTPException(status_code=500, detail="PDF 변환 결과 파일을 찾을 수 없습니다.")
    return pdf_path


# =========================================================
# F-02 API: 통합신청서 자동 작성 + 필요서류 안내문 병합 다운로드
# (이 파일 안에 generate_pdf 함수는 반드시 이 하나만 있어야 합니다)
# =========================================================
@app.post("/api/v1/document/generate")
def generate_pdf(user: UserProfile):
    if not os.path.exists(UNIVERSAL_TEMPLATE_PATH):
        raise HTTPException(status_code=500, detail=f"공통 템플릿 파일이 없습니다: {UNIVERSAL_TEMPLATE_PATH}")

    doc = docx.Document(UNIVERSAL_TEMPLATE_PATH)
    table = doc.tables[0]

    # 성/이름
    r, c = FIELD_MAP["surname"]; set_cell_text_preserving_format(table, r, c, user.surname)
    r, c = FIELD_MAP["given_name"]; set_cell_text_preserving_format(table, r, c, user.given_name)

    # 생년월일
    y, m, d = user.date_of_birth.split("-")
    set_cell_text_preserving_format(table, *FIELD_MAP["birth_year"], y)
    set_cell_text_preserving_format(table, *FIELD_MAP["birth_month"], m)
    set_cell_text_preserving_format(table, *FIELD_MAP["birth_day"], d)

    # 성별
    sex_cfg = FIELD_MAP["sex_male"] if user.sex == "M" else FIELD_MAP["sex_female"]
    set_radio_checkbox_preserving_format(
        table, sex_cfg["row"], sex_cfg["col"], sex_cfg["para"], sex_cfg["target"], sex_cfg["new"]
    )

    # 국적
    nat = FIELD_MAP["nationality"]
    set_paragraph_text_preserving_format(table, nat["row"], nat["col"], nat["para"], user.nationality)

    # 외국인등록번호 (콤필드 13칸)
    digits = list(user.arc_number_front) + list(user.arc_number_back)
    comb_cols = FIELD_MAP["arc_comb_cols"]
    if len(digits) != len(comb_cols):
        raise HTTPException(status_code=500, detail="외국인등록번호 자릿수와 서식 칸 수가 일치하지 않습니다.")
    for col, digit in zip(comb_cols, digits):
        set_cell_text_preserving_format(table, FIELD_MAP["arc_comb_row"], col, digit)

    # 여권
    r, c = FIELD_MAP["passport_number"]; set_cell_text_preserving_format(table, r, c, user.passport_number)
    r, c = FIELD_MAP["passport_issue_date"]; set_cell_text_preserving_format(table, r, c, user.passport_issue_date)
    r, c = FIELD_MAP["passport_expiry_date"]; set_cell_text_preserving_format(table, r, c, user.passport_expiry_date)

    # 주소/연락처
    r, c = FIELD_MAP["address"]; set_cell_text_preserving_format(table, r, c, user.address)
    r, c = FIELD_MAP["phone"]; set_cell_text_preserving_format(table, r, c, user.phone)
    r, c = FIELD_MAP["mobile_phone"]; set_cell_text_preserving_format(table, r, c, user.mobile_phone)
    r, c = FIELD_MAP["home_country_address"]; set_cell_text_preserving_format(table, r, c, user.home_country_address)
    r, c = FIELD_MAP["home_country_phone"]; set_cell_text_preserving_format(table, r, c, user.home_country_phone)
    r, c = FIELD_MAP["email"]; set_cell_text_preserving_format(table, r, c, user.email)

    # 신청일 (자동, 오늘 날짜)
    r, c = FIELD_MAP["application_date"]
    set_cell_text_preserving_format(table, r, c, date.today().strftime("%Y-%m-%d"))

    # 신청구분 체크박스
    cb = APPLICATION_TYPE_MAP[user.application_type]
    set_checkbox_preserving_format(table, cb["row"], cb["col"], cb["target"], cb["new"])

    # 희망 자격 (해당 신청구분인 경우만)
    if "desired_status_cell" in cb:
        ds_row, ds_col = cb["desired_status_cell"]
        set_desired_status_preserving_format(table, ds_row, ds_col, user.desired_status)

    # 필요서류 안내문을 같은 문서 안에 끼워넣기 (LibreOffice 변환 1회로 절감)
    checklist = get_required_documents(user.visa_type, user.application_type)
    prepend_checklist_page(doc, table, checklist)

    # 임시 작업 디렉토리 (원본 템플릿은 절대 덮어쓰지 않음)
    work_dir = f"/tmp/pdfgen_{uuid.uuid4().hex}"
    os.makedirs(work_dir, exist_ok=True)

    try:
        filled_docx_path = os.path.join(work_dir, "filled.docx")
        doc.save(filled_docx_path)
        final_pdf_path = convert_docx_to_pdf(filled_docx_path, work_dir)  # 변환 1번만 실행
    except Exception:
        # 변환 도중 에러가 나도 개인정보가 담긴 임시파일은 즉시 삭제
        shutil.rmtree(work_dir, ignore_errors=True)
        raise

    # 응답(파일 다운로드)이 클라이언트에게 완전히 전송된 직후 임시폴더 삭제
    cleanup_task = BackgroundTask(shutil.rmtree, work_dir, ignore_errors=True)
    return FileResponse(
        final_pdf_path,
        media_type="application/pdf",
        filename="visa_application_package.pdf",
        background=cleanup_task,
    )


SENDGRID_API_KEY = os.environ.get("SENDGRID_API_KEY")
SENDER_EMAIL = os.environ.get("SENDER_EMAIL")  # SendGrid에 인증한 이메일
CRON_SECRET = os.environ.get("CRON_SECRET")     # 외부에서 함부로 못 부르게 막는 비밀키

def send_notification_email(to_email: str, visa_expiry_date: str, application_type: str):
    sg = sendgrid.SendGridAPIClient(api_key=SENDGRID_API_KEY)
    message = Mail(
        from_email=SENDER_EMAIL,
        to_emails=to_email,
        subject="[비자 만료 알림] 체류기간이 30일 이내로 다가왔습니다",
        html_content=f"""
        <p>안녕하세요,</p>
        <p>등록하신 체류자격({application_type})의 만료일이 <b>{visa_expiry_date}</b>로,
        30일 이내로 다가왔습니다.</p>
        <p>지금 서비스에 접속하여 신청서를 미리 준비해보세요.</p>
        """,
    )
    sg.send(message)


@app.post("/api/v1/cron/notify")
def cron_notify(x_cron_secret: str = Header(None)):
    """Cloud Scheduler가 매일 호출. 비밀키 헤더 없이는 실행 안 됨."""
    if not CRON_SECRET or x_cron_secret != CRON_SECRET:
        raise HTTPException(status_code=403, detail="인증되지 않은 요청입니다.")

    target_date = (date.today() + timedelta(days=30)).strftime("%Y-%m-%d")

    docs = fs_client.collection(NOTIFICATION_COLLECTION) \
        .where("visa_expiry_date", "==", target_date).stream()

    sent_count = 0
    for doc in docs:
        data = doc.to_dict()
        send_notification_email(data["email"], data["visa_expiry_date"], data["application_type"])
        sent_count += 1

    return {"status": "success", "notified": sent_count, "target_date": target_date}


class ChatMessage(BaseModel):
    role: str      # "user" 또는 "assistant"
    content: str


class ChatRequest(BaseModel):
    message: str
    history: list[ChatMessage] = []


@app.post("/api/v1/chat")
def chat(req: ChatRequest):
    """서류/신청 절차 관련 질문에 RAG로 검색된 매뉴얼 근거만으로 답변하는 챗봇."""
    relevant_chunks = retrieve_chunks_for_query(req.message, top_k=4)
    context_text = "\n".join(relevant_chunks) if relevant_chunks else "(관련 규정을 찾지 못했습니다)"

    history_text = "\n".join(
        f"{'사용자' if h.role == 'user' else '챗봇'}: {h.content}"
        for h in req.history[-6:]  # 너무 길어지지 않도록 최근 6턴만
    )

    prompt = f"""
    너는 대한민국 외국인 유학생의 체류허가 신청서 작성과 서류 준비를 돕는 안내 챗봇이야.

    아래 [검색된 관련 규정]이 있다면 그 내용을 최우선 근거로 삼아 답변해.
    검색된 자료에 없는 내용이라도, 일반적으로 널리 알려진 정보(예: 특정 서류가 무엇인지,
    보통 어디서 발급받는지, 일반적인 절차 등)는 너의 지식을 활용해 친절하게 답변해도 좋아.

    다만 한국 체류·비자 신청의 구체적인 요건, 정확한 서류 목록, 기한처럼
    틀리면 안 되는 내용은 반드시 검색된 자료에 근거가 있을 때만 단정적으로 답하고,
    근거가 없으면 추측하지 말고 "정확한 안내를 위해 하이코리아(국번없이 1345) 또는
    담당 출입국·외국인관서에 문의해 주세요"라고 안내해.

    친절하고 간결하게, 존댓말로 답변해.

    [검색된 관련 규정]:
    {context_text}

    [이전 대화]:
    {history_text}

    [사용자 질문]: {req.message}
    """

    response = model.generate_content(prompt)
    return {"answer": response.text.strip()}